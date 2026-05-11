from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn

class DocxBuilder:
    def __init__(self, template_path):
        self.doc = Document(template_path)
        self.blue_color = RGBColor(0, 170, 204)  # Твой #00AACC
        self.red_color = RGBColor(204, 0, 0)     # Твой #CC0000

    def add_tour_day(self, day_number, date_str, title, places_data):
        """
        Добавляет блок дня: 2 колонки (Фото | Текст)
        """
        # Заголовок дня
        p = self.doc.add_paragraph()
        run = p.add_run(f"День {day_number} ({date_str})")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = self.blue_color

        # Создаем таблицу 1 ряд, 2 колонки
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = False
        col_photos = table.columns[0]
        col_text = table.columns[1]
        col_photos.width = Inches(2.2)
        col_text.width = Inches(4.3)

        cells = table.rows[0].cells

        # Левая колонка: Фото (вставляем из URL или временных файлов)
        # В идеале: скачиваем фото из R2 во временную папку
        # cells[0].add_paragraph().add_run().add_picture(temp_img_path, width=Inches(2))

        # Правая колонка: Текст
        text_cell = cells[1]
        title_p = text_cell.paragraphs[0]
        title_p.add_run(title).bold = True
        
        for place in places_data:
            p = text_cell.add_paragraph(style='List Bullet')
            p.add_run(f"{place['name']}: ").bold = True
            p.add_run(place['description'])

    def save(self, output_path):
        self.doc.save(output_path)